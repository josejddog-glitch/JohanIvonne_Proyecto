"""Ensambla `resolucion_diligenciada.docx` a partir del JSON estructurado de Claude.

Estructura vigente:

    TÍTULO         (negrita, centrado)
    SEGUNDO        (≥2 párrafos, justificado)
    TERCERO        (≥2 párrafos, justificado)
    CUARTO         (intro + subsecciones 4.x, cada subsección con título en negrita
                   y párrafos justificados; las citas literales de artículos van
                   en itálica con sangría izquierda 1cm)

NO hay PRIMERO, NO hay RESUELVE separado, NO hay firma — el abogado las añade
manualmente al revisar.

Esquema esperado del JSON:

    {
      "titulo": "CUN ... del ...",
      "segundo": ["SEGUNDO: ...", "Adicionalmente..."],   # str o list[str]
      "tercero": ["TERCERO: ...", "Posteriormente..."],
      "cuarto": {                                          # str (legacy) o dict
        "intro": "CUARTO: Que, en virtud de los hechos expuestos, ...",
        "subsecciones": [
          {
            "numero": "4.1",
            "titulo": "Pérdida del número celular en prepago",
            "parrafos": ["Observa la Dirección que ...",
                         "\"ARTÍCULO ... \" (Destacado fuera de texto)",
                         "De lo anterior se colige ...",
                         "Aplicando lo anterior al caso ...",
                         "De este modo, se puede concluir ..."]
          }
        ]
      }
    }
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

# Regex que captura marcadores `[VERIFICAR: ...]` (sin permitir corchetes anidados)
VERIFICAR_RE = re.compile(r'\[VERIFICAR:[^\]]*\]')
COLOR_VERIFICAR = RGBColor(0xC0, 0x00, 0x00)  # rojo oscuro tipo Word

# Markdown inline: combinaciones de ***bolditalic***, **bold**, *italic*, __underline__.
# Las combinaciones más específicas (con más marcadores) deben ir primero.
MARKDOWN_INLINE_RE = re.compile(
    r'(\*\*\*__([^_]+)__\*\*\*'        # ***__BIU__***
    r'|__\*\*\*([^*]+)\*\*\*__'        # __***BIU***__ (orden alterno)
    r'|\*\*__([^_]+)__\*\*'            # **__BU__**
    r'|\*\*\*([^*]+)\*\*\*'            # ***BI***
    r'|\*\*([^*]+)\*\*'                # **B**
    r'|__([^_]+)__'                    # __U__
    r'|\*([^*\s][^*]*?)\*)'            # *I*
)


def _parsear_inline_markdown(texto: str) -> list[tuple[str, bool, bool, bool]]:
    """Divide un texto con marcadores inline en segmentos
    `(texto, bold, italic, underline)`.

    Marcadores:
    - `***__x__***` o `__***x***__` -> bold + italic + underline
    - `**__x__**` -> bold + underline
    - `***x***` -> bold + italic
    - `**x**` -> bold
    - `__x__` -> underline (solo)
    - `*x*` -> italic (sólo si no empieza con espacio)
    - Texto fuera de marcadores -> sin formato.
    """
    if not texto:
        return []
    segmentos: list[tuple[str, bool, bool, bool]] = []
    pos = 0
    for m in MARKDOWN_INLINE_RE.finditer(texto):
        ini, fin = m.span()
        if ini > pos:
            segmentos.append((texto[pos:ini], False, False, False))
        if m.group(2) is not None:        # ***__BIU__***
            segmentos.append((m.group(2), True, True, True))
        elif m.group(3) is not None:      # __***BIU***__
            segmentos.append((m.group(3), True, True, True))
        elif m.group(4) is not None:      # **__BU__**
            segmentos.append((m.group(4), True, False, True))
        elif m.group(5) is not None:      # ***BI***
            segmentos.append((m.group(5), True, True, False))
        elif m.group(6) is not None:      # **B**
            segmentos.append((m.group(6), True, False, False))
        elif m.group(7) is not None:      # __U__
            segmentos.append((m.group(7), False, False, True))
        elif m.group(8) is not None:      # *I*
            segmentos.append((m.group(8), False, True, False))
        pos = fin
    if pos < len(texto):
        segmentos.append((texto[pos:], False, False, False))
    return segmentos


def _limpiar_cuerpo(doc) -> None:
    """Elimina todos los párrafos del cuerpo del documento (preserva estilos,
    encabezados, pies de página y márgenes)."""
    body = doc.element.body
    NS_P = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'
    for p in list(body.iter(NS_P)):
        if p.getparent().tag.endswith('}body'):
            body.remove(p)


def _aplicar_run(run, *, bold: bool, italic: bool, font: str, size: int,
                 color: RGBColor | None = None, underline: bool = False) -> None:
    run.bold = bold
    run.italic = italic
    if underline:
        run.underline = True
    run.font.name = font
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def _add_segmento(
    p,
    texto: str,
    *,
    bold_default: bool,
    italic_default: bool,
    font: str,
    size: int,
    color: RGBColor | None = None,
    underline_default: bool = False,
) -> None:
    """Agrega un segmento al párrafo respetando markdown inline (`***`, `**`,
    `*`, `__`). El bold/italic/underline del segmento es OR del default + lo
    del marcador.
    """
    if not texto:
        return
    segmentos = _parsear_inline_markdown(texto)
    if not segmentos:
        run = p.add_run(texto)
        _aplicar_run(
            run,
            bold=bold_default,
            italic=italic_default,
            font=font,
            size=size,
            color=color,
            underline=underline_default,
        )
        return
    for seg_text, seg_bold, seg_italic, seg_underline in segmentos:
        if not seg_text:
            continue
        run = p.add_run(seg_text)
        _aplicar_run(
            run,
            bold=bold_default or seg_bold,
            italic=italic_default or seg_italic,
            font=font,
            size=size,
            color=color,
            underline=underline_default or seg_underline,
        )


def _add_runs_con_verificar(p, texto: str, *, bold: bool, italic: bool,
                            font: str, size: int) -> None:
    """Agrega runs al párrafo dividiendo el texto cuando aparece un marcador
    `[VERIFICAR: ...]`, que se pinta en rojo + negrita para alta visibilidad
    al revisar el borrador. Respeta markdown inline en los tramos no-VERIFICAR.
    """
    pos = 0
    for m in VERIFICAR_RE.finditer(texto):
        ini, fin = m.span()
        if ini > pos:
            _add_segmento(p, texto[pos:ini], bold_default=bold, italic_default=italic, font=font, size=size)
        # marcador VERIFICAR siempre rojo + negrita, sin parsear markdown adentro
        run = p.add_run(texto[ini:fin])
        _aplicar_run(run, bold=True, italic=italic, font=font, size=size, color=COLOR_VERIFICAR)
        pos = fin
    if pos < len(texto):
        _add_segmento(p, texto[pos:], bold_default=bold, italic_default=italic, font=font, size=size)


def _parrafo(doc, texto: str = "", *, bold: bool = False, italic: bool = False,
             align: str = "justify", size: int = 11, font: str = "Arial",
             indent_cm: float | None = None,
             right_indent_cm: float | None = None):
    p = doc.add_paragraph()
    if align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if indent_cm is not None:
        p.paragraph_format.left_indent = Cm(indent_cm)
    if right_indent_cm is not None:
        p.paragraph_format.right_indent = Cm(right_indent_cm)
    if not texto:
        return p
    if "[VERIFICAR" in texto:
        _add_runs_con_verificar(p, texto, bold=bold, italic=italic, font=font, size=size)
    else:
        _add_segmento(p, texto, bold_default=bold, italic_default=italic, font=font, size=size)
    return p


def _es_cita_literal(texto: str) -> bool:
    """Heurística: el párrafo es parte de una cita literal de un artículo si:
    - empieza con comillas dobles (`"`, `“`, `«`), o
    - contiene la coletilla "(Destacado fuera de texto)", o
    - es un encabezado estructural del manual (`CAPÍTULO N`, `SECCIÓN N`,
      `RÉGIMEN DE PROTECCIÓN...`, `TÍTULO N`) — estos aparecen como párrafos
      sueltos dentro de una cita y deben heredar su formato (italic + sangría),
    - es solo `(…)` de elisión.
    """
    s = texto.strip()
    if not s:
        return False
    if s.startswith(('"', '“', '«')):
        return True
    if "(Destacado fuera de texto)" in s:
        return True
    if _es_heading_estructural(s):
        return True
    # Solo elisión: `(…)`, `(...)`, etc. (sin marcadores markdown alrededor)
    limpio = re.sub(r'\*+|__', '', s).strip()
    if limpio in ("(…)", "(...)", "(. . .)"):
        return True
    return False


# Encabezados estructurales del manual SIC que van CENTRADOS dentro de la cita
# literal. Detección por inicio del texto (sin marcadores markdown). Tolerante
# a la ausencia de tildes (CAPITULO/CAPÍTULO, SECCION/SECCIÓN, etc.).
HEADING_ESTRUCTURAL_RE = re.compile(
    r'^\s*(CAPITULO|REGIMEN DE PROTECCION|SECCION\s+\d+\.?|TITULO\s+[IVX]+\.?)\b',
    re.IGNORECASE,
)


def _quitar_tildes(s: str) -> str:
    """Normaliza eliminando tildes para que CAPÍTULO == CAPITULO."""
    import unicodedata
    return "".join(
        c for c in unicodedata.normalize("NFD", s)
        if unicodedata.category(c) != "Mn"
    )


def _es_heading_estructural(texto: str) -> bool:
    """Devuelve True si el párrafo es un encabezado del manual que va centrado
    (CAPÍTULO N, SECCIÓN N, RÉGIMEN DE PROTECCIÓN..., TÍTULO N). Tolerante
    a tildes, marcadores markdown y comillas/apóstrofes iniciales (que pueden
    aparecer cuando el encabezado está dentro de una cita anidada del manual,
    p.ej. `'***CAPÍTULO 1***`).
    """
    # Quitar marcadores markdown
    limpio = re.sub(r'\*+|__', '', texto).strip()
    # Quitar comillas/apóstrofes iniciales (citas anidadas del manual)
    limpio = re.sub(r'^[\'"‘’“”«»\s]+', '', limpio)
    return bool(HEADING_ESTRUCTURAL_RE.match(_quitar_tildes(limpio)))


def _normalizar_a_lista(valor: Any) -> list[str]:
    """Acepta None / str / list[str] y devuelve list[str] sin vacíos."""
    if valor is None:
        return []
    if isinstance(valor, str):
        # split por doble salto de línea para permitir múltiples párrafos en un solo string
        return [p.strip() for p in valor.split("\n\n") if p.strip()]
    if isinstance(valor, list):
        return [str(x).strip() for x in valor if str(x).strip()]
    return [str(valor).strip()]


def _pintar_parrafos_seccion(doc, parrafos: list[str]) -> None:
    """Pinta una lista de párrafos justificados (uso para SEGUNDO y TERCERO)."""
    for p in parrafos:
        _parrafo(doc, p)


def _pintar_cuarto_estructurado(doc, cuarto: dict) -> None:
    """Pinta el CUARTO con intro + subsecciones (estructura del manual SIC).

    Trackea "cita activa": cuando un párrafo empieza con `"`, todos los
    párrafos siguientes heredan el formato de cita (italic + sangría izq+der
    + justify) hasta que un párrafo termine con `"` o contenga "(Destacado
    fuera de texto)". Así, los párrafos intermedios del cuerpo del artículo
    (que no empiezan ni terminan con comilla) también se renderizan como
    parte de la cita.
    """
    intro = cuarto.get("intro", "").strip()
    if intro:
        _parrafo(doc, intro)

    for sub in cuarto.get("subsecciones", []) or []:
        if not isinstance(sub, dict):
            continue
        _parrafo(doc)  # línea en blanco antes del título
        numero = (sub.get("numero") or "").strip()
        titulo = (sub.get("titulo") or "").strip()
        encabezado = f"{numero}.   {titulo}" if numero and titulo else (numero or titulo)
        if encabezado:
            _parrafo(doc, encabezado, bold=True, align="left")

        # Aplanar elementos que vienen con saltos `\n\n` dentro como múltiples
        # párrafos lógicos.
        parrafos_planos: list[str] = []
        for parr in sub.get("parrafos", []) or []:
            if not isinstance(parr, str) or not parr.strip():
                continue
            for sub_parr in parr.split("\n\n"):
                s = sub_parr.strip()
                if s:
                    parrafos_planos.append(s)

        cita_activa = False
        for s in parrafos_planos:
            empieza_cita = s.startswith(('"', '“', '«'))
            # La cita termina cuando el párrafo termina con comilla o cuando
            # contiene la coletilla "(Destacado fuera de texto)" (que va al
            # final del bloque citado).
            termina_cita = (
                s.rstrip().endswith(('"', '”', '»'))
                or "(Destacado fuera de texto)" in s
            )

            if empieza_cita:
                cita_activa = True

            es_heading = _es_heading_estructural(s)
            limpio = re.sub(r'\*+|__', '', s).strip()
            es_elision = limpio in ("(…)", "(...)", "(. . .)")

            es_parte_cita = cita_activa or es_heading or es_elision

            if es_parte_cita:
                # Citas literales: itálica + sangría izquierda Y derecha 1.5cm.
                # CAPÍTULO/SECCIÓN/RÉGIMEN además van con align=center.
                if es_heading:
                    _parrafo(doc, s, italic=True, align="center",
                             indent_cm=1.5, right_indent_cm=1.5)
                else:
                    _parrafo(doc, s, italic=True,
                             indent_cm=1.5, right_indent_cm=1.5)
            else:
                _parrafo(doc, s)

            # Cerrar el modo cita cuando este párrafo la termina.
            if cita_activa and termina_cita:
                cita_activa = False


def construir(plantilla_path: Path, salida_path: Path, datos: dict[str, Any]) -> Path:
    """Construye y guarda el .docx. Retorna la ruta."""
    salida_path.parent.mkdir(parents=True, exist_ok=True)
    import shutil
    shutil.copyfile(plantilla_path, salida_path)
    doc = Document(str(salida_path))
    _limpiar_cuerpo(doc)

    # Etiqueta de sentido (opcional, sobre el título): CONFIRMA / REVOCA / MODIFICA / IMPROCEDENTE
    etiqueta = str(datos.get("etiqueta_sentido") or "").strip().upper()
    if etiqueta in {"CONFIRMA", "REVOCA", "MODIFICA", "IMPROCEDENTE"}:
        _parrafo(doc, etiqueta, bold=True, align="center", size=12)
        _parrafo(doc)

    # TÍTULO
    titulo_lista = _normalizar_a_lista(datos.get("titulo"))
    titulo = titulo_lista[0] if titulo_lista else ""
    if titulo:
        _parrafo(doc, titulo, bold=True, align="center", size=12)
        _parrafo(doc)

    # SEGUNDO
    segundo = _normalizar_a_lista(datos.get("segundo"))
    if segundo:
        _pintar_parrafos_seccion(doc, segundo)
        _parrafo(doc)

    # TERCERO
    tercero = _normalizar_a_lista(datos.get("tercero"))
    if tercero:
        _pintar_parrafos_seccion(doc, tercero)
        _parrafo(doc)

    # CUARTO: acepta dict (estructurado) o str/list (legacy compacto)
    cuarto = datos.get("cuarto")
    if isinstance(cuarto, dict):
        _pintar_cuarto_estructurado(doc, cuarto)
    else:
        cuarto_parrafos = _normalizar_a_lista(cuarto)
        if cuarto_parrafos:
            _pintar_parrafos_seccion(doc, cuarto_parrafos)

    doc.save(str(salida_path))
    return salida_path
